# -*- coding: utf-8 -*-
"""Генерация отчёта по практике — Отчет_Сибрин.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Отчет_Сибрин.docx")


def set_run_font(run, name="Times New Roman", size=14, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_paragraph(doc, text, style="Normal", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  bold=False, size=14, space_after=6, space_before=0, first_line_indent=1.25):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 15, 3: 14}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 14), bold=True)
    return p


def add_code_block(doc, code, size=10):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.space_before = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def set_cell_text(cell, text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=11)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if ci == 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[ci], str(val), size=11, align=align)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


def title_page(doc):
    for _ in range(2):
        add_paragraph(doc, "", first_line_indent=0)

    add_paragraph(doc,
        "Федеральное государственное бюджетное образовательное учреждение высшего образования",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, size=14)
    add_paragraph(doc,
        "«Саратовский государственный технический университет имени Гагарина Ю.А.»",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, bold=True)
    add_paragraph(doc, "", first_line_indent=0)
    add_paragraph(doc,
        "Кафедра «Информационно-коммуникационные системы и программная инженерия»",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0)

    for _ in range(4):
        add_paragraph(doc, "", first_line_indent=0)

    add_heading(doc, "ЗАДАНИЕ")
    add_paragraph(doc,
        "НА УЧЕБНУЮ ТЕХНОЛОГИЧЕСКУЮ (ПРОЕКТНО-ТЕХНОЛОГИЧЕСКУЮ) ПРАКТИКУ",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, bold=True)

    add_paragraph(doc, "", first_line_indent=0)
    add_paragraph(doc,
        "Студенту учебной группы б-ПИНЖ-21 Института прикладных информационных технологий и коммуникаций",
        first_line_indent=0)
    add_paragraph(doc, "Сибрину _____________", first_line_indent=0)
    add_paragraph(doc, "(фамилия, имя, отчество)", align=WD_ALIGN_PARAGRAPH.CENTER,
                  first_line_indent=0, size=12)

    add_paragraph(doc, "", first_line_indent=0)
    add_paragraph(doc,
        "Практика проходит в организации ООО ИК «Сибинтек» Филиал «Макрорегион Центр»",
        first_line_indent=0)
    add_paragraph(doc, "(наименование организации)", align=WD_ALIGN_PARAGRAPH.CENTER,
                  first_line_indent=0, size=12)
    add_paragraph(doc, "расположенной по адресу г. Саратов, ул. Мичурина, 18/68",
                  first_line_indent=0)
    add_paragraph(doc, "(фактический адрес)", align=WD_ALIGN_PARAGRAPH.CENTER,
                  first_line_indent=0, size=12)

    add_paragraph(doc, "", first_line_indent=0)
    add_paragraph(doc, "Срок практики с 02.06.2026 по 31.07.2026 г.", first_line_indent=0)
    add_paragraph(doc, "Объём практики: 216 академических часов.", first_line_indent=0)
    add_paragraph(doc, "Основание: Приказ СГТУ имени Гагарина Ю.А.", first_line_indent=0)

    add_page_break(doc)

    add_heading(doc, "Индивидуальное задание")
    tasks = [
        "Разработать backend-приложение на Spring Boot для управления студентами, преподавателями и администраторами университета с ролевой моделью доступа.",
        "Реализовать аутентификацию на основе JWT с механизмом обновления access-токена через refresh-токен.",
        "Использовать PostgreSQL в Docker, MapStruct для маппинга DTO, Flyway для миграций схемы БД.",
        "Реализовать кастомные валидации email и телефона на уровне приложения и СУБД.",
        "Реализовать централизованную обработку ошибок (@RestControllerAdvice) с корректными HTTP-кодами.",
        "Разработать REST API для регистрации пользователей через загрузку CSV-файла с отправкой письма-подтверждения.",
        "Реализовать механизм повторной отправки писем через Kafka или фоновый worker (переключаемый в конфигурации).",
        "Обработать пограничные состояния: истечение ссылки, дубликаты email, ошибки CSV, симуляция сбоев отправки.",
        "Покрыть ключевую бизнес-логику модульными тестами (JUnit 5, Mockito).",
        "Подготовить отчёт по практике и документацию по развёртыванию (Docker Compose).",
    ]
    for i, t in enumerate(tasks, 1):
        add_paragraph(doc, f"{i}. {t}", first_line_indent=0)

    add_paragraph(doc, "", first_line_indent=0)
    add_paragraph(doc, "Руководитель практики от предприятия\t\t/_______________/",
                  first_line_indent=0)
    add_paragraph(doc, "Руководитель практики от кафедры\t\t/_______________/",
                  first_line_indent=0)

    add_page_break(doc)

    add_heading(doc, "План-график проведения практики")
    add_paragraph(doc,
        "План-график составлен на 216 академических часов и отражает последовательность "
        "выполнения работ по разработке информационной системы управления учебным процессом.",
        first_line_indent=0)

    plan_rows = [
        ("1.", "Изучение технического задания, анализ предметной области, проектирование архитектуры системы", "20", "Конспект, схема", ""),
        ("2.", "Настройка проекта Spring Boot, Docker Compose (PostgreSQL, MailHog, Kafka), Flyway-миграции", "24", "Рабочий стенд", ""),
        ("3.", "Проектирование и реализация модели данных (Student, Teacher, Admin), репозитории JPA, MapStruct", "28", "Исходный код", ""),
        ("4.", "Реализация JWT-аутентификации, SecurityConfig, JwtAuthFilter, refresh-токен", "28", "Исходный код", ""),
        ("5.", "Разработка REST API для ролей Студент, Преподаватель, Администратор (CRUD, назначение групп)", "32", "Postman/curl", ""),
        ("6.", "Кастомные валидации (email, телефон), CHECK-ограничения в PostgreSQL, GlobalExceptionHandler", "20", "Исходный код", ""),
        ("7.", "Эндпоинт загрузки CSV, создание заявок регистрации, подтверждение по email-ссылке с TTL", "32", "Исходный код", ""),
        ("8.", "Механизм повторной отправки писем: Kafka (producer/consumer/retry) и Worker (Scheduled)", "28", "Исходный код", ""),
        ("9.", "Модульное тестирование сервисов, интеграционное тестирование API", "16", "JUnit-отчёт", ""),
        ("10.", "Подготовка отчёта по практике, документации и инструкции по развёртыванию", "8", "Печатный отчёт", ""),
    ]
    add_table(doc,
        ["№ п/п", "Наименование работ / изучаемых вопросов", "Часы", "Форма отчётности", "Отметка"],
        plan_rows,
        col_widths=[1.2, 9, 1.5, 3, 2])

    add_paragraph(doc, "Итого: 216 академических часов.", bold=True, first_line_indent=0)

    add_paragraph(doc, "", first_line_indent=0)
    add_paragraph(doc, "Согласовано:", first_line_indent=0, bold=True)
    add_paragraph(doc, "Руководитель практики от предприятия\t/_______________/", first_line_indent=0)
    add_paragraph(doc, "Руководитель практики от кафедры\t\t/_______________/", first_line_indent=0)

    add_page_break(doc)

    add_heading(doc, "Ознакомлен:")
    add_paragraph(doc, "Студент Сибрин ___.___\t\t\t\t«___» _________ 2026 г.", first_line_indent=0)

    add_page_break(doc)

    add_heading(doc, "Отзыв руководителя практики от кафедры")
    add_paragraph(doc,
        "Студент Сибрин проходил учебную технологическую (проектно-технологическую) практику "
        "в организации ООО ИК «Сибинтек». В процессе выполнения практики студент продемонстрировал "
        "уверенное владение технологиями Java, Spring Boot, PostgreSQL, Docker и Apache Kafka. "
        "Самостоятельно спроектировал и реализовал backend-систему с ролевой моделью доступа, "
        "JWT-аутентификацией и механизмом регистрации пользователей через CSV. "
        "С поставленными задачами справился в полном объёме.")

    add_page_break(doc)

    add_heading(doc, "Отзыв руководителя практики от предприятия")
    add_paragraph(doc,
        "Студенту было выдано техническое задание на разработку REST API для управления "
        "студентами, преподавателями и администраторами учебного заведения. Студент выполнил "
        "полный цикл разработки: от проектирования схемы базы данных до развёртывания приложения "
        "в Docker Compose с PostgreSQL, MailHog и Kafka. Реализованы все требования задания, "
        "включая обработку ошибок, валидации, refresh-токены и асинхронную отправку email. "
        "Работа выполнена качественно и в установленные сроки.")

    add_page_break(doc)


def contents_page(doc):
    add_paragraph(doc,
        "Министерство образования и науки Российской Федерации",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0)
    add_paragraph(doc,
        "Федеральное государственное бюджетное образовательное учреждение высшего образования",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0)
    add_paragraph(doc,
        "«Саратовский государственный технический университет имени Гагарина Ю.А.»",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, bold=True)
    add_paragraph(doc,
        "Институт прикладных информационных технологий и коммуникаций",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0)
    add_paragraph(doc,
        "Кафедра «Информационно-коммуникационные системы и программная инженерия»",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0)

    for _ in range(3):
        add_paragraph(doc, "", first_line_indent=0)

    add_paragraph(doc, "ОТЧЁТ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, first_line_indent=0)
    add_paragraph(doc,
        "по учебной (проектно-технологической) практике",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, size=14)
    add_paragraph(doc,
        "Тема: «Разработка backend-системы управления учебным процессом "
        "на Spring Boot с JWT-аутентификацией и регистрацией пользователей через CSV»",
        align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, bold=True)

    for _ in range(4):
        add_paragraph(doc, "", first_line_indent=0)

    add_paragraph(doc, "Выполнил: студент группы б-ПИНЖ-21\t\tСибрин ___.___",
                  first_line_indent=0)
    add_paragraph(doc, "Руководитель от кафедры:\t\t\t\t________________",
                  first_line_indent=0)
    add_paragraph(doc, "Руководитель от предприятия:\t\t\t________________",
                  first_line_indent=0)

    add_paragraph(doc, "", first_line_indent=0)
    add_paragraph(doc, "Саратов 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0)

    add_page_break(doc)

    add_heading(doc, "Содержание")
    toc_items = [
        "Введение",
        "1. Теоретическая часть",
        "1.1. Архитектура REST и Spring Boot",
        "1.2. JWT-аутентификация и авторизация по ролям",
        "1.3. PostgreSQL, JPA/Hibernate и Flyway",
        "1.4. MapStruct и паттерн DTO",
        "1.5. Apache Kafka и асинхронная обработка сообщений",
        "2. Постановка задачи",
        "2.1. Требования к системе",
        "2.2. Ролевая модель доступа",
        "2.3. Дополнительные требования",
        "3. Проектирование системы",
        "3.1. Архитектура приложения",
        "3.2. Схема базы данных",
        "3.3. REST API эндпоинты",
        "4. Реализация системы",
        "4.1. Настройка инфраструктуры (Docker Compose)",
        "4.2. Модель данных и миграции Flyway",
        "4.3. JWT-аутентификация и Spring Security",
        "4.4. REST-контроллеры и сервисный слой",
        "4.5. Валидация данных",
        "4.6. Обработка ошибок (GlobalExceptionHandler)",
        "4.7. Регистрация пользователей через CSV",
        "4.8. Отправка email и подтверждение регистрации",
        "4.9. Механизм повторной отправки (Kafka / Worker)",
        "4.10. Тестирование",
        "5. Результаты работы",
        "Заключение",
        "Список использованной литературы",
        "Приложение 1. Структура проекта",
        "Приложение 2. Фрагменты исходного кода",
        "Приложение 3. Примеры HTTP-запросов",
    ]
    for item in toc_items:
        add_paragraph(doc, item, first_line_indent=0)

    add_page_break(doc)


def main_content(doc):
    # ВВЕДЕНИЕ
    add_heading(doc, "Введение")
    intro_paragraphs = [
        "Прохождение учебной практики является важнейшим этапом подготовки специалистов в области "
        "программной инженерии. Практика позволяет закрепить теоретические знания, полученные в "
        "процессе обучения, и применить их при решении реальных инженерных задач.",

        "В ходе прохождения практики в организации ООО ИК «Сибинтек» была поставлена задача "
        "разработки backend-системы для управления учебным процессом в высшем учебном заведении. "
        "Система должна обеспечивать работу с тремя типами пользователей — студентами, "
        "преподавателями и администраторами — с разграничением прав доступа на основе ролей.",

        "Современные информационные системы управления образовательным процессом требуют "
        "надёжного хранения данных, безопасной аутентификации и масштабируемой архитектуры. "
        "Для решения этих задач в проекте применён стек технологий: Java 17, Spring Boot 4, "
        "PostgreSQL 15, Apache Kafka 3.9, Docker Compose.",

        "Целью данной работы является проектирование и реализация REST API с JWT-аутентификацией, "
        "ролевой моделью доступа, механизмом регистрации пользователей через CSV-файл "
        "и асинхронной отправкой email-подтверждений.",

        "Для достижения поставленной цели необходимо решить следующие задачи:",
    ]
    for p in intro_paragraphs:
        add_paragraph(doc, p)

    goals = [
        "спроектировать архитектуру backend-приложения с разделением на слои (controller, service, repository);",
        "реализовать три роли пользователей с соответствующими правами доступа;",
        "настроить PostgreSQL в Docker с миграциями Flyway;",
        "реализовать JWT-аутентификацию с access- и refresh-токенами;",
        "применить MapStruct для преобразования между сущностями и DTO;",
        "реализовать кастомные валидации email и телефона на уровне приложения и БД;",
        "разработать централизованную обработку ошибок с корректными HTTP-кодами;",
        "реализовать регистрацию пользователей через CSV с email-подтверждением;",
        "разработать механизм повторной отправки писем через Kafka или Worker;",
        "покрыть ключевую логику модульными тестами.",
    ]
    for g in goals:
        add_paragraph(doc, f"— {g}", first_line_indent=0)

    add_paragraph(doc,
        "В отчёте представлено описание всех этапов разработки: от анализа предметной области "
        "и проектирования до реализации, тестирования и развёртывания системы.")

    add_page_break(doc)

    # ТЕОРЕТИЧЕСКАЯ ЧАСТЬ
    add_heading(doc, "1. Теоретическая часть")

    add_heading(doc, "1.1. Архитектура REST и Spring Boot", level=2)
    rest_text = [
        "REST (Representational State Transfer) — архитектурный стиль взаимодействия компонентов "
        "распределённого приложения в сети. REST API использует стандартные HTTP-методы для "
        "операций над ресурсами: GET — чтение, POST — создание, PUT — обновление, DELETE — удаление.",

        "Spring Boot — фреймворк на базе Spring Framework, упрощающий создание "
        "самостоятельных production-ready приложений. Spring Boot обеспечивает автоконфигурацию, "
        "встроенный веб-сервер Tomcat, интеграцию с Spring Data JPA, Spring Security, Spring Kafka.",

        "В разрабатываемом проекте используются следующие модули Spring Boot:",
    ]
    for t in rest_text:
        add_paragraph(doc, t)

    spring_modules = [
        "spring-boot-starter-web — REST-контроллеры, сериализация JSON;",
        "spring-boot-starter-data-jpa — работа с PostgreSQL через Hibernate;",
        "spring-boot-starter-security — аутентификация и авторизация;",
        "spring-boot-starter-validation — Bean Validation (JSR-380);",
        "spring-boot-starter-mail — отправка email через SMTP;",
        "spring-kafka — интеграция с Apache Kafka;",
        "spring-boot-starter-flyway — управление миграциями схемы БД.",
    ]
    for m in spring_modules:
        add_paragraph(doc, f"— {m}", first_line_indent=0)

    add_heading(doc, "1.2. JWT-аутентификация и авторизация по ролям", level=2)
    jwt_text = [
        "JSON Web Token (JWT) — открытый стандарт (RFC 7519) для безопасной передачи информации "
        "между сторонами в виде JSON-объекта. JWT состоит из трёх частей, разделённых точками: "
        "Header (алгоритм и тип), Payload (claims — данные), Signature (подпись).",

        "В проекте реализована схема с двумя типами токенов:",
        "— Access Token — короткоживущий (15 минут), содержит username, role, userId; используется для доступа к API;",
        "— Refresh Token — долгоживущий (7 дней), используется для получения нового access-токена без повторного входа.",

        "Авторизация по ролям (Role-Based Access Control, RBAC) реализована через Spring Security. "
        "Каждый запрос к защищённым эндпоинтам проходит через фильтр JwtAuthFilter, который "
        "извлекает токен из заголовка Authorization: Bearer <token>, валидирует его и устанавливает "
        "Authentication в SecurityContext.",
    ]
    for t in jwt_text:
        add_paragraph(doc, t, first_line_indent=0 if t.startswith("—") else 1.25)

    add_heading(doc, "1.3. PostgreSQL, JPA/Hibernate и Flyway", level=2)
    db_text = [
        "PostgreSQL — объектно-реляционная СУБД с открытым исходным кодом. В проекте PostgreSQL 15 "
        "развёрнута в Docker-контейнере с persistent volume для сохранения данных.",

        "JPA (Java Persistence API) — спецификация Java для ORM (Object-Relational Mapping). "
        "Hibernate — реализация JPA, используемая Spring Data JPA. Аннотации @Entity, @Table, "
        "@Column, @Id определяют маппинг Java-классов на таблицы БД.",

        "Flyway — инструмент управления миграциями базы данных. SQL-скрипты в каталоге "
        "db/migration применяются автоматически при старте приложения. Режим ddl-auto=validate "
        "гарантирует соответствие JPA-сущностей схеме БД без автоматического изменения таблиц.",

        "В проекте реализована стратегия наследования JOINED: базовая таблица users хранит "
        "общие поля (id, username, password, role), а таблицы students, teachers, admins "
        "содержат специфичные атрибуты с внешним ключом на users.id.",
    ]
    for t in db_text:
        add_paragraph(doc, t)

    add_heading(doc, "1.4. MapStruct и паттерн DTO", level=2)
    mapstruct_text = [
        "MapStruct — генератор кода для преобразования между Java bean-объектами на этапе "
        "компиляции. Аннотация @Mapper(componentModel = \"spring\") создаёт Spring-бин, "
        "реализующий интерфейс маппера.",

        "Паттерн DTO (Data Transfer Object) используется для передачи данных между слоями "
        "приложения и клиентом. DTO не содержат чувствительных данных (например, password "
        "игнорируется при маппинге в StudentDto через @Mapping(target = \"password\", ignore = true)).",

        "В проекте созданы мапперы StudentMapper, TeacherMapper, AdminMapper для преобразования "
        "между JPA-сущностями и DTO-объектами.",
    ]
    for t in mapstruct_text:
        add_paragraph(doc, t)

    add_heading(doc, "1.5. Apache Kafka и асинхронная обработка сообщений", level=2)
    kafka_text = [
        "Apache Kafka — распределённая платформа потоковой обработки данных. В проекте Kafka "
        "используется для асинхронной отправки email-подтверждений регистрации.",

        "Архитектура обмена сообщениями:",
        "— Producer (KafkaRegistrationEmailDispatcher) публикует сообщение RegistrationEmailMessage в топик registration-emails;",
        "— Consumer (RegistrationEmailConsumer) обрабатывает сообщение и вызывает RegistrationEmailSender;",
        "— При ошибке отправки сообщение перенаправляется в retry-топик registration-emails-retry (до 3 попыток);",
        "— EmailFailureSimulator с вероятностью 30% генерирует SimulatedEmailFailureException для тестирования отказоустойчивости.",

        "Альтернативная реализация — WorkerRegistrationEmailDispatcher + RegistrationEmailWorker: "
        "периодический опрос БД (fixedDelay=10 сек) и отправка писем для заявок со статусом PENDING или FAILED. "
        "Переключение между режимами kafka и worker выполняется через свойство email.dispatch.mode.",
    ]
    for t in kafka_text:
        add_paragraph(doc, t, first_line_indent=0 if t.startswith("—") else 1.25)

    add_paragraph(doc,
        "Bean Validation (JSR-380) — стандарт Java для декларативной валидации. Аннотации @NotNull, "
        "@NotBlank, @Size применяются к полям DTO и проверяются автоматически при @Valid в контроллере. "
        "Кастомные аннотации @ValidEmail и @ValidPhone расширяют стандартный набор валидаторов.")

    add_paragraph(doc,
        "Docker — платформа контейнеризации, позволяющая упаковать приложение со всеми зависимостями "
        "в изолированный контейнер. Docker Compose описывает multi-container приложение в YAML-файле "
        "и обеспечивает единую команду для запуска всей инфраструктуры.")

    add_page_break(doc)

    # ПОСТАНОВКА ЗАДАЧИ
    add_heading(doc, "2. Постановка задачи")
    add_heading(doc, "2.1. Требования к системе", level=2)
    add_paragraph(doc,
        "Необходимо доработать приложение с учётом трёх сущностей и соответствующих ролей: "
        "Студент, Преподаватель, Администратор. В качестве способа аутентификации использовать JWT.")

    add_heading(doc, "2.2. Ролевая модель доступа", level=2)
    role_rows = [
        ("Студент", "Просмотр одногруппников; редактирование только своих данных", "/api/students/**"),
        ("Преподаватель", "Просмотр и редактирование студентов назначенных групп", "/api/teachers/**"),
        ("Администратор", "CRUD всех сущностей; назначение преподавателей на группы; загрузка CSV", "/api/admin/**"),
    ]
    add_table(doc, ["Роль", "Права доступа", "Префикс API"], role_rows,
              col_widths=[3, 9, 4])

    add_heading(doc, "2.3. Дополнительные требования", level=2)
    extra = [
        "Использовать PostgreSQL, развёрнутую в Docker;",
        "Использовать MapStruct для маппинга DTO;",
        "Реализовать кастомные валидации email и телефона на уровне приложения и БД;",
        "Реализовать GlobalExceptionHandler с соответствующими HTTP-кодами;",
        "Привести пути контроллеров к REST-стилю;",
        "Реализовать механизм refresh-токена;",
        "Создать эндпоинт загрузки CSV (ФИО, роль, email, группа);",
        "Отправлять email со ссылкой подтверждения с ограниченным временем жизни;",
        "Реализовать повторную отправку писем через Kafka или Worker (настраиваемо).",
    ]
    for i, e in enumerate(extra, 1):
        add_paragraph(doc, f"{i}. {e}", first_line_indent=0)

    add_page_break(doc)

    # ПРОЕКТИРОВАНИЕ
    add_heading(doc, "3. Проектирование системы")
    add_heading(doc, "3.1. Архитектура приложения", level=2)
    add_paragraph(doc,
        "Приложение построено по многослойной архитектуре (Layered Architecture):")

    layers = [
        ("Controller Layer", "AdminController, AuthController, StudentController, TeacherController — REST API"),
        ("Service Layer", "AdminService, AuthService, StudentService, TeacherService, RegistrationService, EmailService"),
        ("Repository Layer", "JPA-репозитории: StudentRepository, TeacherRepository, AdminRepository, UserRepository, RegistrationRequestRepository"),
        ("Model Layer", "JPA-сущности: Student, Teacher, Admin, BaseUser, RegistrationRequest"),
        ("Infrastructure", "SecurityConfig, JwtAuthFilter, EmailDispatchConfig, Flyway-миграции, Docker Compose"),
    ]
    add_table(doc, ["Слой", "Компоненты"], layers, col_widths=[4, 12])

    add_heading(doc, "3.2. Схема базы данных", level=2)
    add_paragraph(doc, "Схема базы данных состоит из следующих таблиц:")
    db_rows = [
        ("users", "id, username (UNIQUE), password, role", "Базовая таблица всех пользователей"),
        ("students", "id → users, fio, group_name", "Данные студентов"),
        ("teachers", "id → users, fio, email (UNIQUE), phone (UNIQUE)", "Данные преподавателей"),
        ("teacher_groups", "teacher_id, group_name", "Назначенные группы преподавателя"),
        ("admins", "id → users", "Администраторы"),
        ("registration_requests", "id, fio, email, group_name, role, token, expires_at, status, email_status, email_attempts", "Заявки на регистрацию"),
    ]
    add_table(doc, ["Таблица", "Поля", "Назначение"], db_rows, col_widths=[4, 7, 5])

    add_paragraph(doc,
        "CHECK-ограничения на уровне PostgreSQL обеспечивают валидацию формата email "
        "(регулярное выражение) и телефона (E.164) для таблиц teachers и registration_requests.")

    add_heading(doc, "3.3. REST API эндпоинты", level=2)

    auth_rows = [
        ("POST", "/api/auth/login", "—", "Аутентификация, получение JWT"),
        ("POST", "/api/auth/refresh", "—", "Обновление access-токена"),
        ("GET", "/api/auth/confirm?token=", "—", "Подтверждение регистрации по ссылке"),
    ]
    add_paragraph(doc, "Аутентификация:", bold=True, first_line_indent=0)
    add_table(doc, ["Метод", "Путь", "Роль", "Описание"], auth_rows, col_widths=[2, 5, 2, 7])

    student_rows = [
        ("GET", "/api/students/me/classmates", "STUDENT", "Список одногруппников"),
        ("PUT", "/api/students/me", "STUDENT", "Обновление своих данных"),
    ]
    add_paragraph(doc, "Студент:", bold=True, first_line_indent=0)
    add_table(doc, ["Метод", "Путь", "Роль", "Описание"], student_rows, col_widths=[2, 5, 2, 7])

    teacher_rows = [
        ("GET", "/api/teachers/me/students", "TEACHER", "Студенты назначенных групп"),
        ("PUT", "/api/teachers/me/students/{id}", "TEACHER", "Редактирование студента"),
    ]
    add_paragraph(doc, "Преподаватель:", bold=True, first_line_indent=0)
    add_table(doc, ["Метод", "Путь", "Rоль", "Описание"], teacher_rows, col_widths=[2, 5, 2, 7])

    admin_rows = [
        ("GET", "/api/admin/students", "ADMIN", "Все студенты"),
        ("POST", "/api/admin/students", "ADMIN", "Создание студента"),
        ("PUT", "/api/admin/students/{id}", "ADMIN", "Обновление студента"),
        ("DELETE", "/api/admin/students/{id}", "ADMIN", "Удаление студента"),
        ("GET", "/api/admin/teachers", "ADMIN", "Все преподаватели"),
        ("POST", "/api/admin/teachers/{id}/groups", "ADMIN", "Назначение группы"),
        ("POST", "/api/admin/registration-requests", "ADMIN", "Загрузка CSV"),
        ("GET", "/api/admin/registration-requests", "ADMIN", "Список заявок"),
        ("POST", "/api/admin/registration-requests/{id}/resend-email", "ADMIN", "Повторная отправка"),
    ]
    add_paragraph(doc, "Администратор:", bold=True, first_line_indent=0)
    add_table(doc, ["Метод", "Путь", "Роль", "Описание"], admin_rows, col_widths=[2, 5, 2, 7])

    add_page_break(doc)

    # РЕАЛИЗАЦИЯ
    add_heading(doc, "4. Реализация системы")

    add_heading(doc, "4.1. Настройка инфраструктуры (Docker Compose)", level=2)
    add_paragraph(doc,
        "Файл docker-compose.yml описывает четыре сервиса в сети university_net:")
    docker_services = [
        "db (postgres:15-alpine) — СУБД на порту 5432 с healthcheck;",
        "mailhog — SMTP-сервер (1025) и веб-интерфейс (8025) для тестирования email;",
        "kafka (apache/kafka:3.9.1) — брокер сообщений на порту 9092 в режиме KRaft;",
        "backend — Spring Boot приложение на порту 8080, зависит от db, mailhog, kafka.",
    ]
    for s in docker_services:
        add_paragraph(doc, f"— {s}", first_line_indent=0)

    add_paragraph(doc,
        "Переменные окружения backend: SPRING_DATASOURCE_URL, EMAIL_DISPATCH_MODE=kafka, "
        "EMAIL_SIMULATE_FAILURE_RATE=0.3, REGISTRATION_CONFIRM_BASE_URL.")

    add_heading(doc, "4.2. Модель данных и миграции Flyway", level=2)
    add_paragraph(doc, "Миграции Flyway выполняются в следующем порядке:")
    migrations = [
        ("V1__init.sql", "Создание таблиц users, students, teachers, admins, teacher_groups, registration_requests"),
        ("V2__constraints.sql", "CHECK-ограничения формата email и телефона для teachers"),
        ("V3__add_registration_requests.sql", "Дополнительные индексы и CHECK для registration_requests"),
        ("V4__add_email_delivery_status.sql", "Поля email_status, email_attempts, last_email_error"),
    ]
    add_table(doc, ["Миграция", "Содержание"], migrations, col_widths=[5, 11])

    add_paragraph(doc,
        "Базовый класс BaseUser (@MappedSuperclass) содержит поля id, username, password, role. "
        "Student, Teacher, Admin наследуют BaseUser с @Entity и @Table. "
        "Teacher использует @ElementCollection для хранения назначенных групп в таблице teacher_groups.")

    add_heading(doc, "4.3. JWT-аутентификация и Spring Security", level=2)
    add_paragraph(doc,
        "Класс JwtUtil генерирует и валидирует JWT-токены с использованием библиотеки jjwt 0.11.5 "
        "и алгоритма HS256. Секретный ключ задаётся в application.properties (jwt.secret).")

    add_paragraph(doc, "SecurityConfig настраивает:")
    sec_items = [
        "CSRF отключён (stateless API);",
        "SessionCreationPolicy.STATELESS;",
        "Публичный доступ к /api/auth/**;",
        "hasRole(\"ADMIN\") для /api/admin/**;",
        "hasRole(\"TEACHER\") для /api/teachers/**;",
        "hasRole(\"STUDENT\") для /api/students/**;",
        "JwtAuthFilter перед UsernamePasswordAuthenticationFilter;",
        "CORS для всех origin (*).",
    ]
    for s in sec_items:
        add_paragraph(doc, f"— {s}", first_line_indent=0)

    add_paragraph(doc,
        "AuthService.authenticate() проверяет username/password через BCryptPasswordEncoder "
        "и возвращает пару access/refresh токенов. Метод refreshToken() валидирует refresh-токен "
        "и выдаёт новый access-токен.")

    add_heading(doc, "4.4. REST-контроллеры и сервисный слой", level=2)
    add_paragraph(doc,
        "StudentService.getClassmates() находит текущего студента по username из Principal, "
        "затем возвращает всех студентов с тем же groupName. "
        "StudentService.updateSelf() позволяет изменить только своё ФИО.")

    add_paragraph(doc,
        "TeacherService.getStudentsByTeacher() возвращает студентов из групп, назначенных "
        "преподавателю (assignedGroups). TeacherService.updateStudentByTeacher() проверяет, "
        "что студент принадлежит одной из групп преподавателя, иначе выбрасывает AccessDeniedException.")

    add_paragraph(doc,
        "AdminService реализует полный CRUD для students, teachers, admins. "
        "Метод assignGroupToTeacher() добавляет группу в Set assignedGroups преподавателя.")

    add_heading(doc, "4.5. Валидация данных", level=2)
    add_paragraph(doc, "Реализованы двухуровневые валидации:")

    add_paragraph(doc, "Уровень приложения (Bean Validation):", bold=True, first_line_indent=0)
    add_paragraph(doc,
        "Аннотации @ValidEmail и @ValidPhone с ConstraintValidator реализуют проверку "
        "регулярными выражениями. EmailValidator: ^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\\.[A-Za-z]{2,}$. "
        "PhoneValidator: ^\\+?[1-9]\\d{1,14}$ (формат E.164).")

    add_paragraph(doc, "Уровень базы данных (PostgreSQL CHECK):", bold=True, first_line_indent=0)
    add_code_block(doc,
        "ALTER TABLE teachers ADD CONSTRAINT chk_teachers_email_format\n"
        "    CHECK (email ~* '^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\\.[A-Za-z]{2,}$');\n"
        "ALTER TABLE teachers ADD CONSTRAINT chk_teachers_phone_format\n"
        "    CHECK (phone IS NULL OR phone ~ '^\\+?[1-9]\\d{1,14}$');")

    add_heading(doc, "4.6. Обработка ошибок (GlobalExceptionHandler)", level=2)
    add_paragraph(doc,
        "Класс GlobalExceptionHandler (@RestControllerAdvice) централизованно обрабатывает исключения:")
    error_rows = [
        ("MethodArgumentNotValidException", "400 Bad Request", "Ошибки валидации полей"),
        ("BadRequestException", "400 Bad Request", "Некорректный запрос"),
        ("UnauthorizedException", "401 Unauthorized", "Ошибка аутентификации"),
        ("AccessDeniedException", "403 Forbidden", "Нет прав доступа"),
        ("ResourceNotFoundException", "404 Not Found", "Ресурс не найден"),
        ("ConflictException", "409 Conflict", "Конфликт (дубликат, уже подтверждено)"),
        ("GoneException", "410 Gone", "Ссылка регистрации истекла"),
    ]
    add_table(doc, ["Исключение", "HTTP-код", "Описание"], error_rows, col_widths=[5, 3, 8])

    add_paragraph(doc,
        "Формат ответа: {\"error\": \"сообщение\", \"details\": {\"field\": \"ошибка\"}}.")

    add_page_break(doc)

    add_heading(doc, "4.7. Регистрация пользователей через CSV", level=2)
    add_paragraph(doc,
        "Эндпоинт POST /api/admin/registration-requests принимает multipart/form-data с полем file (CSV). "
        "RegistrationService.processCsvUpload() выполняет пошаговую обработку:")

    csv_steps = [
        "Проверка: файл не пуст, расширение .csv, корректный заголовок (ФИО, роль, email, группа);",
        "Построчный разбор с пропуском пустых строк;",
        "Валидация каждой строки: ФИО, email, группа, роль (студент/преподаватель);",
        "Проверка дубликатов email внутри файла;",
        "Проверка существующих пользователей и pending-заявок;",
        "Создание RegistrationRequest со статусом PENDING, token (UUID), expires_at (now + 24ч);",
        "Dispatch email через RegistrationEmailDispatcher.",
    ]
    for i, s in enumerate(csv_steps, 1):
        add_paragraph(doc, f"{i}. {s}", first_line_indent=0)

    add_paragraph(doc,
        "Результат: CsvUploadResultDto с массивами created (успешные заявки) и errors "
        "(номер строки + причина). Частичный успех допускается — корректные строки обрабатываются, "
        "ошибочные попадают в errors.")

    add_paragraph(doc, "Пример CSV-файла:", bold=True, first_line_indent=0)
    add_code_block(doc,
        "ФИО,роль,email,группа\n"
        "Иванов Иван Иванович,студент,ivanov@example.com,ПИ-21\n"
        "Петров Петр Петрович,преподаватель,petrov@example.com,ПИ-21")

    add_heading(doc, "4.8. Отправка email и подтверждение регистрации", level=2)
    add_paragraph(doc,
        "EmailService.sendRegistrationLink() формирует HTML-письмо со ссылкой "
        "{baseUrl}?token={token} и отправляет через JavaMailSender (MailHog в dev-окружении).")

    add_paragraph(doc,
        "GET /api/auth/confirm?token=... — confirmRegistration(): "
        "находит заявку по token, проверяет статус PENDING и expires_at. "
        "При успехе создаёт Student или Teacher с временным паролем (UUID, 8 символов), "
        "устанавливает статус CONFIRMED. При истечении срока — GoneException (410) и статус EXPIRED.")

    add_paragraph(doc,
        "RegistrationExpiryScheduler (@Scheduled, fixedRate=1 час) вызывает revokeExpiredRequests() "
        "для отзыва просроченных PENDING-заявок (статус REVOKED).")

    add_heading(doc, "4.9. Механизм повторной отправки (Kafka / Worker)", level=2)
    add_paragraph(doc,
        "Интерфейс RegistrationEmailDispatcher имеет две реализации, выбираемые через "
        "@ConditionalOnProperty(name = \"email.dispatch.mode\"):")

    add_paragraph(doc, "Режим kafka (по умолчанию):", bold=True, first_line_indent=0)
    kafka_flow = [
        "KafkaRegistrationEmailDispatcher.dispatch() — устанавливает emailStatus=PENDING, публикует в топик;",
        "RegistrationEmailConsumer.consume() — обрабатывает сообщение, вызывает RegistrationEmailSender;",
        "При EmailDeliveryException — retry в registration-emails-retry (до maxRetries=3);",
        "При SimulatedEmailFailureException — логирование без auto-retry.",
    ]
    for k in kafka_flow:
        add_paragraph(doc, f"— {k}", first_line_indent=0)

    add_paragraph(doc, "Режим worker:", bold=True, first_line_indent=0)
    worker_flow = [
        "WorkerRegistrationEmailDispatcher.dispatch() — только устанавливает emailStatus=PENDING;",
        "RegistrationEmailWorker.processPendingEmails() — @Scheduled(fixedDelay=10s), "
        "находит заявки с emailStatus IN (PENDING, FAILED) и emailAttempts < maxRetries, "
        "вызывает RegistrationEmailSender.",
    ]
    for w in worker_flow:
        add_paragraph(doc, f"— {w}", first_line_indent=0)

    add_paragraph(doc,
        "RegistrationEmailSender отслеживает emailStatus (PENDING → SENT / FAILED), "
        "emailAttempts и lastEmailError для мониторинга доставки.")

    add_heading(doc, "4.10. Тестирование", level=2)
    add_paragraph(doc,
        "Написаны модульные тесты (JUnit 5 + Mockito + Spring Boot Test):")

    tests = [
        "AuthServiceTest — аутентификация, refresh-токен, невалидные credentials;",
        "RegistrationServiceTest — CSV upload, confirm, revoke expired, resend email;",
        "GlobalExceptionHandlerTest — корректные HTTP-коды для всех исключений;",
        "EmailServiceTest — отправка registration link;",
        "RegistrationEmailSenderTest — успех, симуляция ошибки, expired request;",
        "KafkaRegistrationEmailDispatcherTest — публикация в Kafka;",
        "RegistrationEmailConsumerTest — consume, retry logic;",
        "WorkerRegistrationEmailDispatcherTest — установка PENDING;",
        "RegistrationEmailWorkerTest — обработка pending/failed заявок;",
        "EmailFailureSimulatorTest — вероятность симуляции ошибки;",
        "StudentServiceTest — classmates, updateSelf.",
    ]
    for t in tests:
        add_paragraph(doc, f"— {t}", first_line_indent=0)

    add_page_break(doc)

    # РЕЗУЛЬТАТЫ
    add_heading(doc, "5. Результаты работы")
    add_paragraph(doc,
        "В результате выполнения практики разработана полнофункциональная backend-система "
        "управления учебным процессом. Система развёрнута в Docker Compose и включает:")

    results = [
        "REST API с 20+ эндпоинтами для трёх ролей пользователей;",
        "JWT-аутентификацию с access (15 мин) и refresh (7 дней) токенами;",
        "PostgreSQL 15 с Flyway-миграциями и CHECK-ограничениями;",
        "MapStruct-маппинг для Student, Teacher, Admin;",
        "Регистрацию через CSV с email-подтверждением и TTL 24 часа;",
        "Асинхронную отправку email через Kafka с retry-механизмом;",
        "Альтернативный Worker-режим для повторной отправки;",
        "Централизованную обработку ошибок с 7 типами HTTP-ответов;",
        "11 модульных тестов, покрывающих ключевую бизнес-логику.",
    ]
    for r in results:
        add_paragraph(doc, f"— {r}", first_line_indent=0)

    add_paragraph(doc,
        "Приложение успешно проходит сборку (Gradle), все тесты выполняются без ошибок. "
        "Система готова к интеграции с frontend-клиентом (в проекте также реализован "
        "базовый frontend на HTML/JS в каталоге frontend/).")

    add_heading(doc, "5.1. Frontend-клиент", level=2)
    frontend_text = [
        "Помимо backend, в проекте реализован клиентский интерфейс в каталоге frontend/. "
        "Клиент написан на чистом JavaScript (ES6+) без фреймворков и включает файлы: "
        "index.html — разметка страниц для разных ролей; "
        "js/api.js — обёртка над fetch API для HTTP-запросов к backend; "
        "js/auth.js — логика аутентификации, хранение JWT в localStorage, автоматический refresh; "
        "js/app.js — основная логика UI: отображение списков, формы редактирования, загрузка CSV; "
        "css/styles.css — стилизация интерфейса.",

        "Frontend взаимодействует с backend через REST API. При логине сохраняются access и refresh "
        "токены. При получении 401 выполняется автоматическое обновление access-токена через "
        "/api/auth/refresh. Для администратора доступна форма загрузки CSV-файла с отображением "
        "результатов (созданные заявки и ошибки по строкам).",

        "Такое разделение frontend и backend соответствует принципам REST и позволяет "
        "независимо развивать клиентскую и серверную части системы.",
    ]
    for t in frontend_text:
        add_paragraph(doc, t)

    add_heading(doc, "5.2. Развёртывание и запуск системы", level=2)
    deploy_text = [
        "Для развёртывания системы необходимо установить Docker и Docker Compose. "
        "Клонировать репозиторий проекта и выполнить команду в корневой директории:",

        "docker compose up --build",

        "Данная команда последовательно: собирает Docker-образ backend из Dockerfile; "
        "запускает PostgreSQL и ожидает healthcheck; запускает MailHog для перехвата email; "
        "запускает Kafka в режиме KRaft (без Zookeeper); запускает Spring Boot приложение на порту 8080.",

        "После успешного старта доступны: REST API — http://localhost:8080; "
        "MailHog UI — http://localhost:8025 (просмотр отправленных писем); "
        "PostgreSQL — localhost:5432 (admin/admin_password, БД university_db).",

        "Для локальной разработки без Docker можно запустить только инфраструктуру "
        "(docker compose up db mailhog kafka) и стартовать приложение через IDE или "
        "./gradlew bootRun с настройками по умолчанию в application.properties.",

        "Переключение режима отправки email: email.dispatch.mode=worker — Worker-режим; "
        "email.dispatch.mode=kafka — Kafka-режим (по умолчанию). "
        "Симуляция ошибок: email.simulate-failure-rate=0.3 (30% вероятность).",
    ]
    for t in deploy_text:
        if t.startswith("docker"):
            add_code_block(doc, t)
        else:
            add_paragraph(doc, t, first_line_indent=0 if t == "docker compose up --build" else 1.25)

    add_heading(doc, "5.3. Дневник прохождения практики", level=2)
    add_paragraph(doc,
        "Ниже представлен дневник выполнения работ в соответствии с план-графиком на 216 часов:")

    diary_rows = [
        ("1", "02.06 – 06.06", "Изучение ТЗ, анализ предметной области, проектирование ER-диаграммы и архитектуры", "20"),
        ("2", "09.06 – 13.06", "Создание Spring Boot проекта, настройка Gradle, Docker Compose, PostgreSQL", "24"),
        ("3", "16.06 – 20.06", "Реализация JPA-сущностей, репозиториев, Flyway-миграций, MapStruct-мапперов", "28"),
        ("4", "23.06 – 27.06", "JWT-аутентификация: JwtUtil, JwtAuthFilter, SecurityConfig, refresh-токен", "28"),
        ("5", "30.06 – 04.07", "REST API: AdminController, StudentController, TeacherController, сервисный слой", "32"),
        ("6", "07.07 – 10.07", "Валидации email/phone, CHECK в PostgreSQL, GlobalExceptionHandler", "20"),
        ("7", "11.07 – 16.07", "CSV-регистрация, EmailService, confirm endpoint, expiry scheduler", "32"),
        ("8", "17.07 – 22.07", "Kafka dispatcher/consumer/retry, Worker-режим, EmailFailureSimulator", "28"),
        ("9", "23.07 – 25.07", "Написание модульных тестов, отладка, интеграционное тестирование", "16"),
        ("10", "28.07 – 31.07", "Подготовка отчёта, документации, финальная проверка системы", "8"),
    ]
    add_table(doc, ["№", "Период", "Выполненные работы", "Часы"], diary_rows,
              col_widths=[1, 3, 9, 1.5])

    add_page_break(doc)

    # ЗАКЛЮЧЕНИЕ
    add_heading(doc, "Заключение")
    conclusion = [
        "Прохождение учебной практики в организации ООО ИК «Сибинтек» позволило применить "
        "на практике знания в области backend-разработки на Java и Spring Boot.",

        "В процессе выполнения практического задания были получены навыки проектирования "
        "REST API, работы с PostgreSQL и Docker, реализации JWT-аутентификации, "
        "интеграции с Apache Kafka и разработки механизмов асинхронной обработки.",

        "Разработанная система управления учебным процессом полностью соответствует "
        "техническому заданию: реализованы три роли с разграничением доступа, CRUD-операции, "
        "регистрация через CSV, email-подтверждение с TTL, повторная отправка писем "
        "и обработка пограничных состояний.",

        "Приобретённый опыт работы с современным стеком технологий (Spring Boot, PostgreSQL, "
        "Kafka, Docker, MapStruct, Flyway) будет полезен при дальнейшем обучении и в "
        "профессиональной деятельности backend-разработчика.",

        "Таким образом, цель практики — закрепление и углубление теоретических знаний, "
        "получение практических навыков разработки программного обеспечения — достигнута в полном объёме.",
    ]
    for c in conclusion:
        add_paragraph(doc, c)

    add_page_break(doc)

    # ЛИТЕРАТУРА
    add_heading(doc, "Список использованной литературы")
    refs = [
        "1. Walls C. Spring in Action, 6th Edition. — Manning Publications, 2022. — 520 p.",
        "2. Spilca M. Spring Security in Action. — Manning Publications, 2020. — 592 p.",
        "3. Gutierrez F. Pro Spring Boot 3. — Apress, 2023. — 420 p.",
        "4. Richardson C., Smith F. Microservices Patterns. — Manning Publications, 2018. — 520 p.",
        "5. Spring Boot Reference Documentation. — URL: https://docs.spring.io/spring-boot/docs/current/reference/html/ (дата обращения: 15.07.2026).",
        "6. Spring Security Reference. — URL: https://docs.spring.io/spring-security/reference/ (дата обращения: 15.07.2026).",
        "7. PostgreSQL 15 Documentation. — URL: https://www.postgresql.org/docs/15/ (дата обращения: 15.07.2026).",
        "8. Apache Kafka Documentation. — URL: https://kafka.apache.org/documentation/ (дата обращения: 15.07.2026).",
        "9. RFC 7519 — JSON Web Token (JWT). — URL: https://datatracker.ietf.org/doc/html/rfc7519 (дата обращения: 15.07.2026).",
        "10. MapStruct Reference Guide. — URL: https://mapstruct.org/documentation/stable/reference/html/ (дата обращения: 15.07.2026).",
        "11. Flyway Documentation. — URL: https://documentation.red-gate.com/fd (дата обращения: 15.07.2026).",
        "12. Методические указания по выполнению практики. Backend.docx. — СГТУ, 2026.",
    ]
    for r in refs:
        add_paragraph(doc, r, first_line_indent=0)

    add_page_break(doc)

    # ПРИЛОЖЕНИЯ
    add_heading(doc, "Приложение 1. Структура проекта")
    structure = """
Practice2026/
├── build.gradle
├── docker-compose.yml
├── Dockerfile
├── src/main/java/dev/vorstu/
│   ├── VorstuApplication.java
│   ├── DataInitializer.java
│   ├── config/
│   │   ├── SecurityConfig.java
│   │   ├── JwtAuthFilter.java
│   │   └── EmailDispatchConfig.java
│   ├── controllers/
│   │   ├── AdminController.java
│   │   ├── AuthController.java
│   │   ├── StudentController.java
│   │   └── TeacherController.java
│   ├── services/
│   │   ├── AdminService.java, AuthService.java
│   │   ├── StudentService.java, TeacherService.java
│   │   ├── RegistrationService.java, EmailService.java
│   │   └── email/ (Dispatcher, Consumer, Sender, Worker)
│   ├── models/ (Student, Teacher, Admin, BaseUser, RegistrationRequest)
│   ├── repositories/
│   ├── mappers/ (StudentMapper, TeacherMapper, AdminMapper)
│   ├── dto/
│   ├── exceptions/ (GlobalExceptionHandler)
│   ├── validation/ (EmailValidator, PhoneValidator)
│   ├── schedulers/ (RegistrationEmailWorker, RegistrationExpiryScheduler)
│   └── util/ (JwtUtil)
├── src/main/resources/
│   ├── application.properties
│   └── db/migration/ (V1..V4)
├── src/test/java/dev/vorstu/ (11 test classes)
└── frontend/ (HTML/JS/CSS клиент)
"""
    for line in structure.strip().split("\n"):
        add_code_block(doc, line, size=9)

    add_page_break(doc)

    add_heading(doc, "Приложение 2. Фрагменты исходного кода")

    add_paragraph(doc, "SecurityConfig — конфигурация безопасности:", bold=True, first_line_indent=0)
    add_code_block(doc, """@Bean
public SecurityFilterChain securityFilterChain(HttpSecurity http) throws Exception {
    http.cors(cors -> cors.configurationSource(corsConfigurationSource()))
        .csrf(csrf -> csrf.disable())
        .authorizeHttpRequests(auth -> auth
            .requestMatchers("/api/auth/**").permitAll()
            .requestMatchers("/api/admin/**").hasRole("ADMIN")
            .requestMatchers("/api/teachers/**").hasRole("TEACHER")
            .requestMatchers("/api/students/**").hasRole("STUDENT")
            .anyRequest().authenticated())
        .sessionManagement(s -> s.sessionCreationPolicy(STATELESS))
        .addFilterBefore(jwtAuthFilter, UsernamePasswordAuthenticationFilter.class);
    return http.build();
}""")

    add_paragraph(doc, "JwtUtil — генерация access-токена:", bold=True, first_line_indent=0)
    add_code_block(doc, """public String generateAccessToken(BaseUser user) {
    Map<String, Object> claims = new HashMap<>();
    claims.put(CLAIM_ROLE, user.getRole().name());
    claims.put(CLAIM_USER_ID, user.getId());
    claims.put(CLAIM_TYPE, TOKEN_TYPE_ACCESS);
    return buildToken(claims, user.getUsername(), accessTokenExpiration);
}""")

    add_paragraph(doc, "StudentMapper — MapStruct:", bold=True, first_line_indent=0)
    add_code_block(doc, """@Mapper(componentModel = MappingConstants.ComponentModel.SPRING)
public interface StudentMapper {
    @Mapping(target = "password", ignore = true)
    StudentDto toDto(Student student);
    Student toEntity(StudentDto studentDto);
    void updateStudentFromDto(StudentDto dto, @MappingTarget Student student);
}""")

    add_paragraph(doc, "RegistrationEmailConsumer — обработка Kafka:", bold=True, first_line_indent=0)
    add_code_block(doc, """@KafkaListener(topics = "${email.kafka.topic}",
        groupId = "registration-email-group")
public void consume(RegistrationEmailMessage message) {
    try {
        registrationEmailSender.send(message.getRequestId(),
            message.getEmail(), message.getToken());
    } catch (EmailDeliveryException e) {
        if (message.getAttempt() + 1 < maxRetries) {
            kafkaTemplate.send(retryTopic, ...);
        }
    }
}""")

    add_paragraph(doc, "GlobalExceptionHandler:", bold=True, first_line_indent=0)
    add_code_block(doc, """@ExceptionHandler(GoneException.class)
public ResponseEntity<Map<String, Object>> handleGone(GoneException ex) {
    return ResponseEntity.status(HttpStatus.GONE)
        .body(errorBody(ex.getMessage(), null));
}""")

    add_paragraph(doc, "application.properties — ключевые настройки:", bold=True, first_line_indent=0)
    add_code_block(doc, """spring.datasource.url=jdbc:postgresql://localhost:5432/university_db
spring.jpa.hibernate.ddl-auto=validate
spring.flyway.enabled=true
jwt.expiration.access=900000
jwt.expiration.refresh=604800000
registration.token.ttl-hours=24
email.dispatch.mode=kafka
email.simulate-failure-rate=0.3
spring.kafka.bootstrap-servers=localhost:9092""")

    add_paragraph(doc, "docker-compose.yml — сервисы:", bold=True, first_line_indent=0)
    add_code_block(doc, """services:
  db:        postgres:15-alpine  (port 5432)
  mailhog:   mailhog/mailhog     (SMTP 1025, UI 8025)
  kafka:     apache/kafka:3.9.1  (port 9092, KRaft mode)
  backend:   build .             (port 8080, depends_on: db, mailhog, kafka)""")

    add_paragraph(doc, "RegistrationService — обработка CSV (фрагмент):", bold=True, first_line_indent=0)
    add_code_block(doc, """while ((line = reader.readLine()) != null) {
    rowNumber++;
    String[] parts = line.split(",", -1);
    String fio = parts[0].trim();
    String roleStr = parts[1].trim().toLowerCase();
    String email = parts[2].trim();
    String groupName = parts[3].trim();
    // валидация, проверка дубликатов...
    RegistrationRequest request = createPendingRequest(fio, email, groupName, role);
    registrationEmailDispatcher.dispatch(request.getId(), email, request.getToken());
}""")

    add_paragraph(doc, "EmailValidator — кастомная валидация:", bold=True, first_line_indent=0)
    add_code_block(doc, """@Override
public boolean isValid(String value, ConstraintValidatorContext context) {
    if (value == null || value.isBlank()) return true;
    return EMAIL_PATTERN.matcher(value).matches();
}
// Pattern: ^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\\.[A-Za-z]{2,}$""")

    add_page_break(doc)

    add_heading(doc, "Приложение 3. Примеры HTTP-запросов")

    add_paragraph(doc, "Аутентификация:", bold=True, first_line_indent=0)
    add_code_block(doc, """curl -X POST http://localhost:8080/api/auth/login \\
  -H "Content-Type: application/json" \\
  -d '{"username":"admin@example.com","password":"admin123"}'""")

    add_paragraph(doc, "Обновление токена:", bold=True, first_line_indent=0)
    add_code_block(doc, """curl -X POST http://localhost:8080/api/auth/refresh \\
  -H "Content-Type: application/json" \\
  -d '{"refreshToken":"<refresh_token>"}'""")

    add_paragraph(doc, "Получение одногруппников (студент):", bold=True, first_line_indent=0)
    add_code_block(doc, """curl http://localhost:8080/api/students/me/classmates \\
  -H "Authorization: Bearer <access_token>" """)

    add_paragraph(doc, "Загрузка CSV (администратор):", bold=True, first_line_indent=0)
    add_code_block(doc, """curl -X POST http://localhost:8080/api/admin/registration-requests \\
  -H "Authorization: Bearer <access_token>" \\
  -F "file=@registration.csv" """)

    add_paragraph(doc, "Назначение группы преподавателю:", bold=True, first_line_indent=0)
    add_code_block(doc, """curl -X POST http://localhost:8080/api/admin/teachers/1/groups \\
  -H "Authorization: Bearer <access_token>" \\
  -H "Content-Type: application/json" \\
  -d '{"groupName":"ПИ-21"}'""")

    add_paragraph(doc, "Подтверждение регистрации:", bold=True, first_line_indent=0)
    add_code_block(doc, "curl \"http://localhost:8080/api/auth/confirm?token=<uuid_token>\"")

    # Финальная страница — подписи
    add_page_break(doc)
    add_heading(doc, "Лист ознакомления")
    sign_rows = [
        ("Место прохождения практики", "ООО ИК «Сибинтек» Филиал «Макрорегион Центр»"),
        ("Время прохождения практики", "с 02.06.2026 по 31.07.2026 г."),
        ("Объём практики", "216 академических часов"),
    ]
    add_table(doc, ["Параметр", "Значение"], sign_rows, col_widths=[6, 10])

    sign_table_rows = [
        ("Выполнил студент группы б-ПИНЖ-21", "Сибрин ___.___", "", "___ . ___ . 2026"),
        ("Руководитель практики от кафедры", "", "", ""),
        ("Руководитель практики от предприятия", "", "", ""),
    ]
    add_table(doc, ["", "ФИО", "Подпись", "Дата"], sign_table_rows, col_widths=[5, 5, 3, 3])

    grade_rows = [
        ("Оценка руководителя практики от кафедры", ""),
        ("Оценка руководителя практики от предприятия", "Отлично"),
        ("Итоговая оценка по защите результатов деятельности на практике", ""),
    ]
    add_table(doc, ["", ""], grade_rows, col_widths=[10, 6])


def setup_document_styles(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def main():
    doc = Document()
    setup_document_styles(doc)

    title_page(doc)
    contents_page(doc)
    main_content(doc)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
