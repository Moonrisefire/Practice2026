from docx import Document
import sys

def read_docx(path, out_path):
    doc = Document(path)
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(f'=== {path} ===\n')
        f.write(f'Paragraphs: {len(doc.paragraphs)}\n')
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                style = p.style.name if p.style else 'None'
                f.write(f'[{i}] ({style}): {p.text}\n')
        f.write(f'Tables: {len(doc.tables)}\n')
        for ti, table in enumerate(doc.tables):
            f.write(f'--- Table {ti} ---\n')
            for row in table.rows:
                cells = [c.text.replace('\n',' ') for c in row.cells]
                f.write(' | '.join(cells) + '\n')

if __name__ == '__main__':
    read_docx(sys.argv[1], sys.argv[2])
